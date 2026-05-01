import fitz
from docx import Document
from docx.shared import Pt

pdf_path = r"c:\Users\Tom\Desktop\剧本agent\郑宝珠B.pdf"
output_docx_path = r"c:\Users\Tom\Desktop\剧本agent\郑宝珠B_提取文字.docx"

print("正在处理PDF...")

# 创建Word文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# 打开PDF
pdf_doc = fitz.open(pdf_path)
total_pages = len(pdf_doc)
print(f"PDF共 {total_pages} 页")

for page_num in range(total_pages):
    page = pdf_doc[page_num]
    
    # 尝试获取文本（如果有的话）
    text = page.get_text("text")
    
    if text.strip():
        # 添加页码
        doc.add_heading(f"第 {page_num + 1} 页", level=2)
        doc.add_paragraph(text)
        doc.add_paragraph("\n")
    else:
        # 如果没有文本，说明是纯图片PDF
        # 由于需要OCR，这里先保存提示信息
        doc.add_heading(f"第 {page_num + 1} 页（图片页，需要OCR）", level=2)
        doc.add_paragraph("此页为图片格式，需要OCR（光学字符识别）才能提取文字。")
        doc.add_paragraph("\n")

# 保存Word文档
doc.save(output_docx_path)
print(f"\n文档已保存到: {output_docx_path}")
print("\n注意：此PDF是扫描件（图片格式），普通PDF解析无法提取文字。")
print("要提取这种PDF的文字，需要使用OCR工具。")
