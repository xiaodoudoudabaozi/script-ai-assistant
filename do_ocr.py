import fitz
import easyocr
from docx import Document
from docx.shared import Pt
import io
from PIL import Image
import numpy as np

pdf_path = r"c:\Users\Tom\Desktop\剧本agent\郑宝珠B.pdf"
output_docx_path = r"c:\Users\Tom\Desktop\剧本agent\郑宝珠B_提取文字.docx"

print("正在初始化OCR...")
# 初始化OCR阅读器，支持中文和英文
reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)

print("正在处理PDF...")
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

pdf_doc = fitz.open(pdf_path)
total_pages = len(pdf_doc)
print(f"PDF共 {total_pages} 页")

for page_num in range(total_pages):
    print(f"正在处理第 {page_num + 1} 页...")
    page = pdf_doc[page_num]
    
    # 将PDF页面转换为图片
    pix = page.get_pixmap(dpi=200)
    img_data = pix.tobytes("png")
    img = Image.open(io.BytesIO(img_data))
    
    # 使用OCR识别文字
    results = reader.readtext(np.array(img), detail=0)
    
    # 添加到Word文档
    doc.add_heading(f"第 {page_num + 1} 页", level=2)
    if results:
        for text in results:
            if text.strip():
                doc.add_paragraph(text)
    else:
        doc.add_paragraph("（未识别到文字）")
    doc.add_paragraph("\n")

# 保存文档
doc.save(output_docx_path)
print(f"\n完成！文档已保存到: {output_docx_path}")
